import io
import base64
from docx import Document
from urllib.parse import quote
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app.core.dependencies import require_owner
from app.core.database import get_db
from app.models.user import User
from app.models.contract import Contract
from app.models.room import Room
from app.models.tenant import Tenant
from app.schemas.contract_schema import ContractCreate, ContractUpdate, ContractResponse


router = APIRouter(prefix="/contracts", tags=["contracts"])

# Hàm hỗ trợ duyệt qua các đoạn văn và bảng biểu trong Word để thay thế dữ liệu
def replace_text_in_doc(doc, replacements):
    # Xử lý văn bản thường
    for p in doc.paragraphs:
        _replace_text_in_paragraph(p, replacements)
                
    # Xử lý văn bản bên trong các Bảng (Tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_text_in_paragraph(p, replacements)


def _replace_text_in_paragraph(p, replacements):
    for key, val in replacements.items():
        if key in p.text:
            for run in p.runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(val))
            
            if key in p.text and not any(key in r.text for r in p.runs):
                p.text = p.text.replace(key, str(val))


@router.post("", response_model=ContractResponse, status_code=status.HTTP_201_CREATED)
def create_contract(payload: ContractCreate, db: Session = Depends(get_db)):
    """Tạo hợp đồng mới
    3 bước validate nghiệp vụ phải làm trước khi insert:
    1. room_id tồn tại
    2. tenant_id tồn tại
    3. room đó chưa có hợp đồng active nào khác
    Sau khi tạo thành công, đồng bộ room.status -> 'occupied'."""

    room = db.query(Room).filter(Room.id == payload.room_id).first()
    if room is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Không tìm thấy phòng có id={payload.room_id}"
        )
    
    tenant = db.query(Tenant).filter(Tenant.id == payload.tenant_id).first()
    if tenant is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Không tìm thấy khách thuê có id={payload.tenant_id}"
        )
    
    existing_active = (
        db.query(Contract)
        .filter(Contract.room_id == payload.room_id, Contract.status == "active")
        .first()
    )
    if existing_active is not None:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail=(
                f"Phòng id={payload.room_id} đang có hợp đồng active "
                f"(id={existing_active.id}). Phải kết thúc hợp đồng cũ "
                "trước khi tạo hợp đồng mới cho phòng này."
            )
        )
    
    new_contract = Contract(**payload.model_dump(), status="active")
    db.add(new_contract)

    # Đồng bộ trạng thái phòng - tránh chủ trọ quên đổi status tay
    room.status = "occupied"

    db.commit()
    db.refresh(new_contract)
    return new_contract


@router.get("", response_model=list[ContractResponse])
def list_contracts(db: Session = Depends(get_db)):
    return db.query(Contract).all()


@router.get("/{contract_id}", response_model=ContractResponse)
def get_contract(contract_id: int, db: Session = Depends(get_db)):
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if contract is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Không tìm thấy hợp đồng có id={contract_id}"
        )
    return contract


@router.patch("/{contract_id}", response_model=ContractResponse)
def update_contract(contract_id: int, payload: ContractUpdate, db: Session = Depends(get_db), current_user: User = Depends(require_owner)):
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if contract is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Không tìm thấy hợp đồng có id={contract_id}"
        )
    
    update_data = payload.model_dump(exclude_unset=True)
    new_status = update_data.get("status")

    for field, value in update_data.items():
        setattr(contract, field, value)

    if new_status in ("ended", "terminated"):
        room = db.query(Room).filter(Room.id == contract.room_id).first()
        room.status = "vacant"

    db.commit()
    db.refresh(contract)
    return contract


@router.delete("/{contract_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_contract(contract_id: int, db: Session = Depends(get_db)):
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if contract is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Không tìm thấy hợp đồng có id={contract_id}"
        )
    
    from app.models.bill import Bill
    has_bill = db.query(Bill).filter(Bill.contract_id == contract_id).first() is not None
    if has_bill:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail=f"Không thể xoá hợp đồng này vì còn bill liên kết (lịch sử)."
        )
    
    db.delete(contract)
    db.commit()
    return None


@router.get("/{contract_id}/export-word")
def export_contract_word(
    contract_id: int, 
    electric_price: float = 0, 
    water_price: float = 0,
    electric_reading: float = 0,
    water_reading: float = 0,
    db: Session = Depends(get_db)):
    contract = db.query(Contract).filter(Contract.id == contract_id).first()
    if not contract:
        raise HTTPException(status_code=404, detail="Hợp đồng không tồn tại")
        
    room = contract.room
    house = room.house
    tenant = contract.tenant

    # Tạo từ điển dữ liệu thay thế an toàn
    replacements = {
        "[NGAY_TAO]": datetime.now().strftime("%d/%m/%Y"),
        "[TEN_NHA]": getattr(house, "name", ""),
        "[DIA_CHI_NHA]": getattr(house, "address", ""),
        "[SO_PHONG]": getattr(room, "room_number", ""),
        "[TEN_CHU_NHA]": getattr(house, "owner_name", ""),
        "[TEN_KHACH]": getattr(tenant, "full_name", ""),
        "[SDT_KHACH]": getattr(tenant, "phone", ""),
        
        # Sửa thành id_card_number theo đúng data của frontend
        "[CCCD_KHACH]": getattr(tenant, "id_card_number", ""), 
        
        "[GIA_THUE]": f"{int(contract.monthly_rent):,} VNĐ" if contract.monthly_rent else "0 VNĐ",
        "[TIEN_COC]": f"{int(contract.deposit):,} VNĐ" if contract.deposit else "0 VNĐ",
        
        # Sử dụng params truyền từ React lên
        "[GIA_DIEN]": f"{int(electric_price):,} VNĐ",
        "[GIA_NUOC]": f"{int(water_price):,} VNĐ",
        "[CHI_SO_DIEN]": str(electric_reading),
        "[CHI_SO_NUOC]": str(water_reading),
        
        "[NGAY_BAT_DAU]": contract.start_date.strftime("%d/%m/%Y") if contract.start_date else "...",
        "[NGAY_KET_THUC]": contract.end_date.strftime("%d/%m/%Y") if contract.end_date else "...",
        "[SO_NGUOI]": str(contract.num_tenants) if contract.num_tenants else "...",
        "[SO_XE]": str(contract.num_vehicles) if contract.num_vehicles else "...",
        "[NOI_THAT]": ", ".join(getattr(room, "furnitures", [])) if room and getattr(room, "furnitures", None) else "Không có",
    }

    template_data = getattr(house, 'contract_template', None)
    doc = None

    # Kiểm tra xem dữ liệu lưu trong DB có phải là Base64 của file ZIP/DOCX không (Bắt đầu bằng UEs)
    if template_data and template_data.startswith("UEs"):
        try:
            file_bytes = base64.b64decode(template_data)
            doc = Document(io.BytesIO(file_bytes))
        except Exception:
            raise HTTPException(status_code=500, detail="Mẫu hợp đồng gốc bị lỗi định dạng.")
    else:
        # Fallback nếu chủ trọ chưa Upload file mẫu, tạo một file trắng cơ bản
        doc = Document()
        doc.add_heading(f"HỢP ĐỒNG THUÊ PHÒNG {room.room_number if room else ''}", 0)
        doc.add_paragraph("Chủ trọ chưa cấu hình file Word mẫu. Vui lòng vào Cài đặt Mẫu Hợp đồng để Upload file .docx.")
        for key, val in replacements.items():
            doc.add_paragraph(f"{key}: {val}")

    # Chạy hàm thay thế dữ liệu
    replace_text_in_doc(doc, replacements)
        
    # Đóng gói file trả về client
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    room_number_safe = room.room_number if room else "Unknown"
    encoded_file_name = quote(f"HopDong_Phong_{room_number_safe}.docx")
    
    return StreamingResponse(
        output, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{encoded_file_name}"}
    )